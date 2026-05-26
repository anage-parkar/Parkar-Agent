"""
Generate the Parkar knowledge-base .docx files.

Run:
    python scripts/generate_docs.py

Produces one .docx per knowledge category in data/docx/. Content is freshly
authored representative material for Parkar. Internal HR documents (Leave
Policy, Work Culture, Upcoming Projects) are marked as SAMPLE so HR/Ops can
replace them with the authoritative versions without changing the pipeline.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

# Make `config` importable whether run as a script or a module.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from config import DOCX_DIR  # noqa: E402

BRAND_TEAL = RGBColor(0x00, 0xC4, 0xA7)
NAVY = RGBColor(0x0B, 0x1F, 0x3A)


# ---------------------------------------------------------------------------
# Rendering helpers
# ---------------------------------------------------------------------------
def _render(spec: dict) -> Document:
    doc = Document()

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_heading(spec["title"], level=0)
    for run in title.runs:
        run.font.color.rgb = NAVY

    if spec.get("subtitle"):
        p = doc.add_paragraph()
        run = p.add_run(spec["subtitle"])
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = BRAND_TEAL

    if spec.get("note"):
        p = doc.add_paragraph()
        run = p.add_run(spec["note"])
        run.bold = True
        run.font.size = Pt(9)

    for section in spec["sections"]:
        h = doc.add_heading(section["heading"], level=1)
        for run in h.runs:
            run.font.color.rgb = NAVY

        for para in section.get("paras", []):
            doc.add_paragraph(para)

        for bullet in section.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

        table_spec = section.get("table")
        if table_spec:
            headers = table_spec["headers"]
            rows = table_spec["rows"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            for i, head in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = head
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)

    return doc


def build_all() -> list[str]:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    written = []
    for filename, spec in DOCS.items():
        doc = _render(spec)
        out = DOCX_DIR / f"{filename}.docx"
        doc.save(str(out))
        written.append(out.name)
    return written


# ---------------------------------------------------------------------------
# CONTENT
# ---------------------------------------------------------------------------
SAMPLE_NOTE = (
    "INTERNAL SAMPLE DOCUMENT — representative content for the AI knowledge base. "
    "Replace with the authoritative HR/Operations version before relying on it."
)

DOCS: dict[str, dict] = {
    # ------------------------------------------------------------------ #
    "leave-policy": {
        "title": "Parkar Leave Policy",
        "subtitle": "People & Culture — Time Off and Leave Entitlements",
        "note": SAMPLE_NOTE,
        "sections": [
            {
                "heading": "Overview",
                "paras": [
                    "Parkar's leave policy is designed to help employees rest, recover and balance "
                    "personal commitments while ensuring teams can plan around absences. The policy "
                    "applies to all full-time employees. The leave year runs from 1 January to 31 December.",
                ],
            },
            {
                "heading": "Annual Leave Entitlements",
                "paras": ["Full-time employees accrue paid leave as set out below."],
                "table": {
                    "headers": ["Leave Type", "Days per Year", "Notes"],
                    "rows": [
                        ["Earned / Privilege Leave", "21", "Accrues monthly; up to 10 days may carry over."],
                        ["Casual Leave", "7", "For short, unplanned personal needs."],
                        ["Sick Leave", "10", "Medical certificate required beyond 2 consecutive days."],
                        ["Public Holidays", "12", "Per the annual holiday calendar by location."],
                    ],
                },
            },
            {
                "heading": "Parental Leave",
                "bullets": [
                    "Maternity leave: 26 weeks of paid leave, in line with statutory requirements.",
                    "Paternity leave: 10 working days of paid leave, to be taken within 3 months of birth or adoption.",
                    "Adoption leave: equivalent entitlement to maternity/paternity leave based on primary caregiver status.",
                ],
            },
            {
                "heading": "How to Apply for Leave",
                "bullets": [
                    "Submit leave requests through the HRMS portal at least 3 working days in advance for planned leave.",
                    "Sick or emergency leave should be reported to your manager as early as possible on the first day.",
                    "Manager approval is required; HR is notified automatically once approved.",
                    "Leave balances are visible in the HRMS dashboard at any time.",
                ],
            },
            {
                "heading": "Other Leave",
                "bullets": [
                    "Bereavement leave: up to 5 days for an immediate family member.",
                    "Marriage leave: 3 days, once per employment period.",
                    "Compensatory off: granted for approved work on a public holiday or weekend.",
                    "Unpaid leave: by exception, with manager and HR approval.",
                ],
            },
            {
                "heading": "Contact",
                "paras": [
                    "For any questions about leave, entitlements or balances, contact the People & Culture (HR) "
                    "team via the HR helpdesk or your HR business partner.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "partners": {
        "title": "Parkar Partners & Alliances",
        "subtitle": "Technology partnerships that accelerate enterprise AI delivery",
        "sections": [
            {
                "heading": "Our Partner Ecosystem",
                "paras": [
                    "Parkar works alongside the world's leading cloud, data and AI platform providers. "
                    "These partnerships mean we bring certified expertise, partner-funded programmes and "
                    "early access to new capabilities to every engagement — so clients reach production faster "
                    "and with less risk.",
                ],
            },
            {
                "heading": "Cloud & Hyperscaler Partners",
                "bullets": [
                    "Microsoft Azure — co-engineering on Azure AI, Azure OpenAI, Fabric and AKS; Microsoft solution partner.",
                    "Amazon Web Services (AWS) — cloud migration, data engineering and well-architected delivery.",
                    "Google Cloud — data analytics and AI workloads.",
                ],
            },
            {
                "heading": "Data & AI Platform Partners",
                "bullets": [
                    "Databricks — lakehouse, data engineering and ML platform delivery.",
                    "Snowflake — cloud data warehousing and cost-optimised analytics.",
                    "Microsoft Fabric — unified analytics and real-time intelligence.",
                ],
            },
            {
                "heading": "Why It Matters",
                "paras": [
                    "Because we don't start from scratch on every engagement: our platforms (AIONIQ and Vector) "
                    "plus partner accelerators give enterprises pre-built AI operating models, data connectors, "
                    "AIOps correlation and reference architectures from day one.",
                ],
            },
            {
                "heading": "Becoming a Partner",
                "paras": [
                    "Organisations interested in partnering with Parkar — whether as a technology, delivery or "
                    "go-to-market partner — can reach out through the Contact page on parkar.in.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "upcoming-projects": {
        "title": "Parkar Upcoming Projects & Initiatives",
        "subtitle": "Forward-looking programmes across platforms, delivery and R&D",
        "note": SAMPLE_NOTE,
        "sections": [
            {
                "heading": "Platform Roadmap",
                "bullets": [
                    "AIONIQ v-next: expanded agentic workflow library and a self-serve AI Readiness Assessment.",
                    "Vector 2.x: deeper multi-cloud observability correlation and 1,000+ MCP server integrations.",
                    "Expanded connector catalogue for data onboarding across enterprise SaaS systems.",
                ],
            },
            {
                "heading": "Delivery Programmes (Representative)",
                "table": {
                    "headers": ["Initiative", "Focus Area", "Stage"],
                    "rows": [
                        ["Gen AI Studio rollout", "Agentic AI / GCC", "Scaling"],
                        ["Real-time AI Ops", "IT Operations", "In delivery"],
                        ["Enterprise data fabric", "Data Engineering", "Discovery"],
                        ["Healthcare population-health insights", "Industries / Data", "Pilot"],
                    ],
                },
            },
            {
                "heading": "Research & Innovation",
                "bullets": [
                    "Agent governance and observability: monitoring and accountability for fleets of AI agents.",
                    "Retrieval-augmented generation (RAG) accelerators for regulated industries.",
                    "Cost-optimisation tooling for cloud and data platforms.",
                ],
            },
            {
                "heading": "Get Involved",
                "paras": [
                    "Employees can express interest in upcoming projects through their manager or the staffing/"
                    "resourcing team. Clients interested in early-access programmes can contact their Parkar "
                    "account team.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "teams": {
        "title": "Parkar Teams & Leadership",
        "subtitle": "How Parkar is organised to deliver enterprise AI",
        "sections": [
            {
                "heading": "How We're Organised",
                "paras": [
                    "Parkar is organised around five core practice areas supported by platform engineering, "
                    "delivery, and corporate functions. Cross-functional pods bring data, cloud, applications, "
                    "IT operations and agentic AI specialists together on each engagement.",
                ],
            },
            {
                "heading": "Practice Teams",
                "bullets": [
                    "Agentic AI — designing and shipping autonomous, governed AI agents and copilots.",
                    "Data Engineering — pipelines, lakehouse, data fabric and analytics foundations.",
                    "Cloud — migration, landing zones, modernization and cost optimization.",
                    "Applications — cloud-native, composable and modern application engineering.",
                    "IT Operations (Vector) — AIOps, observability and reliability.",
                ],
            },
            {
                "heading": "Platform & Engineering",
                "paras": [
                    "Dedicated platform teams build and maintain AIONIQ (the AI operating model platform) and "
                    "Vector (intelligent IT operations and observability). These teams turn repeatable delivery "
                    "patterns into product capability.",
                ],
            },
            {
                "heading": "Leadership",
                "paras": [
                    "Parkar's leadership team spans engineering, delivery, product, industry domains and corporate "
                    "functions. Leadership profiles and the latest organisational details are published on the "
                    "Leadership page at parkar.in/company/leadership.",
                ],
            },
            {
                "heading": "Global Delivery",
                "paras": [
                    "Parkar operates a global delivery model with teams supporting clients across time zones, "
                    "including Global Capability Center (GCC) operations in India.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "platforms": {
        "title": "Parkar Platforms — AIONIQ & Vector",
        "subtitle": "Proprietary platforms that make enterprise AI repeatable",
        "sections": [
            {
                "heading": "AIONIQ — AI Operating Model Platform",
                "paras": [
                    "AIONIQ is Parkar's AI operating model platform. It governs the full journey from AI readiness "
                    "through to production agentic workflows, so enterprises don't have to assemble the operating "
                    "model from scratch.",
                ],
                "bullets": [
                    "AI Readiness Assessment — scored, prioritised, actionable output in days.",
                    "Pre-built AI operating model and agentic workflow library.",
                    "Governance, security and compliance built in (SOC 2 certified).",
                    "1,000+ MCP servers for tool and system integration.",
                ],
            },
            {
                "heading": "Vector — Intelligent IT Operations & Observability",
                "paras": [
                    "Vector is Parkar's intelligent IT operations and observability layer. It keeps cloud, "
                    "applications and AI environments reliable and cost-efficient.",
                ],
                "bullets": [
                    "Unified, multi-cloud observability and AIOps correlation.",
                    "Alert noise reduction and root-cause analysis.",
                    "12+ data connectors for fast onboarding.",
                    "Cost and performance optimisation across environments.",
                ],
            },
            {
                "heading": "Vector Platform & Integrations",
                "paras": [
                    "The Vector platform exposes integrations across monitoring, logging, cloud and ITSM systems, "
                    "consolidating signals into a single operational picture. Details are on the Vector pages at "
                    "parkar.in/platforms/vector.",
                ],
            },
            {
                "heading": "Designed to Work Together",
                "paras": [
                    "AIONIQ and Vector are designed to work together but each delivers standalone value. Parkar "
                    "recommends the right fit based on a client's priorities.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "solutions": {
        "title": "Parkar Solutions & Practice Areas",
        "subtitle": "Five integrated practices that ship enterprise AI to production",
        "sections": [
            {
                "heading": "Agentic AI",
                "paras": [
                    "We design and deploy autonomous, purposeful, governed AI agents — from copilots to "
                    "multi-agent systems — that are connected to enterprise systems and accountable in production.",
                ],
            },
            {
                "heading": "Data Engineering",
                "paras": [
                    "Trusted, governed and connected data foundations: pipelines, lakehouse and data fabric/mesh "
                    "architectures on Azure, AWS, Databricks and Snowflake that make data accurate, accessible "
                    "and observable.",
                ],
            },
            {
                "heading": "Cloud",
                "paras": [
                    "Scalable, secure and optimised cloud: migration, landing zones, modernization and cost "
                    "optimization that make environments resilient, portable and multi-cloud ready.",
                ],
            },
            {
                "heading": "Applications",
                "paras": [
                    "Modern, composable, cloud-native application engineering: microservices, low-code where it "
                    "fits, and maintainable, performant, secure software.",
                ],
            },
            {
                "heading": "IT Operations",
                "paras": [
                    "Resilient, automated and proactive operations powered by Vector — observable, responsive and "
                    "reliable, reducing alert fatigue and improving application performance.",
                ],
            },
            {
                "heading": "One Integrated Programme",
                "paras": [
                    "Each practice works as a standalone engagement or as one integrated programme. Many clients "
                    "start with one area and expand — 95% of clients return for the next programme. Clients can "
                    "engage any scope they need.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "legal-policies": {
        "title": "Parkar Legal Policies",
        "subtitle": "Privacy, terms and compliance",
        "note": "Summary for the knowledge base. The authoritative legal text lives at parkar.in/legal.",
        "sections": [
            {
                "heading": "Privacy Policy (Summary)",
                "bullets": [
                    "Parkar collects only the personal information needed to provide services and respond to enquiries.",
                    "Data is processed lawfully, stored securely, and not sold to third parties.",
                    "Users may request access to, correction of, or deletion of their personal data.",
                    "Full details are published at parkar.in/legal/privacy.",
                ],
            },
            {
                "heading": "Terms of Use (Summary)",
                "bullets": [
                    "Website content is provided for general information about Parkar's services.",
                    "Trademarks, platform names (AIONIQ, Vector) and content are the property of Parkar.",
                    "Use of the site is subject to the full terms at parkar.in/legal/terms.",
                ],
            },
            {
                "heading": "Security & Compliance",
                "bullets": [
                    "AIONIQ is SOC 2 certified.",
                    "HIPAA-compliant data governance is available for healthcare engagements.",
                    "Parkar holds ISO 27001-aligned and AICPA SOC credentials (see certification badges on the site).",
                    "Security and governance are built in to delivery, not bolted on.",
                ],
            },
            {
                "heading": "Contact for Legal & Data Requests",
                "paras": [
                    "Privacy, data-subject and legal enquiries can be submitted through the Contact page on "
                    "parkar.in.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "work-culture": {
        "title": "Life at Parkar — Work Culture",
        "subtitle": "How we work, grow and support each other",
        "note": SAMPLE_NOTE,
        "sections": [
            {
                "heading": "Our Culture",
                "paras": [
                    "Parkar is an outcomes-driven, engineering-led company. We value curiosity, ownership and being "
                    "easy to work with. People are trusted to make decisions close to the work, and we invest in "
                    "mentorship, learning and career growth.",
                ],
            },
            {
                "heading": "Ways of Working",
                "bullets": [
                    "Hybrid working with flexibility, anchored by collaborative in-person time.",
                    "Cross-functional pods so engineers learn across data, cloud, apps, IT ops and AI.",
                    "A speak-up culture — ideas and concerns are welcomed at every level.",
                    "Continuous feedback rather than once-a-year reviews.",
                ],
            },
            {
                "heading": "Growth & Learning",
                "bullets": [
                    "Mentorship programmes pairing engineers with senior leaders.",
                    "Certification support across Microsoft Azure, AWS, Databricks and Snowflake.",
                    "Internal tech talks, communities of practice and hands-on platform training.",
                ],
            },
            {
                "heading": "Wellbeing & Belonging",
                "bullets": [
                    "Mental-health awareness and support initiatives.",
                    "Inclusive, respectful environment with a focus on integrating people and culture.",
                    "Recognition programmes celebrating teams and individuals.",
                ],
            },
            {
                "heading": "Learn More",
                "paras": [
                    "Read more about life at Parkar, team highlights and culture stories at "
                    "parkar.in/careers/life-at-parkar.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "industries": {
        "title": "Industries Parkar Serves",
        "subtitle": "Deep domain context where it matters most",
        "sections": [
            {
                "heading": "Why Industry Context Matters",
                "paras": [
                    "Data models, compliance requirements and AI use cases differ fundamentally across sectors. "
                    "Generic approaches rarely ship to production in regulated environments — so Parkar brings "
                    "deep domain context to each industry.",
                ],
            },
            {
                "heading": "Financial Services",
                "paras": [
                    "Risk and compliance automation, fraud monitoring, and solving the customer-data problems "
                    "banks can't engineer their way out of — with the governance regulated finance demands.",
                ],
            },
            {
                "heading": "Healthcare & Life Sciences",
                "paras": [
                    "HIPAA-compliant data governance, population-health insights, and resilient AI infrastructure "
                    "so patient-facing systems stay reliable.",
                ],
            },
            {
                "heading": "Manufacturing & Industrial",
                "paras": [
                    "Predictive maintenance, supply-chain analytics, and RAG/prompt-engineering use cases that "
                    "bring AI to the factory and the supply chain.",
                ],
            },
            {
                "heading": "Media & High-Tech / SaaS",
                "paras": [
                    "Scalable SaaS application engineering and AI-era product development for high-tech and media "
                    "businesses.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "careers": {
        "title": "Careers at Parkar",
        "subtitle": "Build enterprise AI that actually ships to production",
        "sections": [
            {
                "heading": "Why Join Parkar",
                "paras": [
                    "Parkar engineers work on real production AI for enterprise clients — not just demos. With "
                    "proprietary platforms (AIONIQ, Vector), 250+ engagements and 95% client retention, you work "
                    "across data, cloud, applications, IT operations and agentic AI.",
                ],
            },
            {
                "heading": "Where We Hire",
                "bullets": [
                    "Engineering — data, cloud, application, platform and AI/ML engineers.",
                    "Agentic AI and GenAI specialists.",
                    "IT operations / SRE / observability (Vector).",
                    "Delivery, product and industry domain roles.",
                    "Corporate functions — People & Culture, Finance, Marketing, Sales.",
                ],
            },
            {
                "heading": "Hiring Process",
                "bullets": [
                    "Application via the Open Positions page on parkar.in/careers/open-positions.",
                    "Recruiter screen to understand your background and interests.",
                    "Technical / role-specific interviews with the team.",
                    "Final conversation with leadership and an offer.",
                ],
            },
            {
                "heading": "Growth & Team Highlights",
                "paras": [
                    "Parkar invests in mentorship, certifications and internal mobility. Read team highlights and "
                    "employee stories at parkar.in/careers/team-highlights.",
                ],
            },
            {
                "heading": "How to Apply",
                "paras": [
                    "Browse current openings at parkar.in/careers/open-positions and apply directly. For general "
                    "enquiries, use the Contact page.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "gcc": {
        "title": "Parkar GCC — Global Capability Centers",
        "subtitle": "Build, operate and scale your India GCC for AI",
        "sections": [
            {
                "heading": "Overview",
                "paras": [
                    "Parkar helps enterprises stand up and scale Global Capability Centers (GCCs) — particularly "
                    "in India — that are built not just for scale but for intelligence. We bring AIOps and GenAI "
                    "to unlock productivity and performance.",
                ],
            },
            {
                "heading": "Engagement Models",
                "table": {
                    "headers": ["Model", "Best For", "What You Get"],
                    "rows": [
                        ["Build-Operate-Transfer (BOT)", "Standing up a new GCC", "Parkar builds and operates, then transfers ownership to you."],
                        ["Managed Service", "Running capability as a service", "Parkar runs the capability against agreed outcomes and SLAs."],
                        ["Pay-as-you-go", "Flexible, variable demand", "Scale capacity up or down with consumption-based commercials."],
                    ],
                },
            },
            {
                "heading": "GCC + AI",
                "bullets": [
                    "GenAI and AIOps to enhance efficiency and accelerate innovation.",
                    "Gen AI Studio and real-time AI Ops accelerators.",
                    "Thought leadership on what India's GCCs need to win with GenAI.",
                ],
            },
            {
                "heading": "Learn More",
                "paras": [
                    "Explore GCC overview, models, case studies and thought leadership at parkar.in/gcc.",
                ],
            },
        ],
    },
    # ------------------------------------------------------------------ #
    "insights": {
        "title": "Parkar Insights — Blogs, Research, Events & News",
        "subtitle": "Perspectives on enterprise AI, data, cloud and IT operations",
        "sections": [
            {
                "heading": "What You'll Find",
                "paras": [
                    "Parkar Insights brings together blogs, research, case studies, webinars, events and news. "
                    "Topics span enterprise AI strategy, data engineering, cloud, IT operations/AIOps, and the "
                    "realities of getting AI into production.",
                ],
            },
            {
                "heading": "Popular Themes",
                "bullets": [
                    "Why enterprise AI struggles — split accountability, readiness gaps, and the three gaps leaders miss.",
                    "Data engineering — why pipelines work in pilot but fail in production; data fabric vs data mesh.",
                    "AIOps & observability — moving beyond alerting; reducing alert noise; unified monitoring.",
                    "GCCs & GenAI — building for intelligence, not just scale.",
                    "Cloud — landing zones, cost optimization, and cloud-native security by design.",
                ],
            },
            {
                "heading": "Case Studies",
                "bullets": [
                    "Predictive maintenance in manufacturing.",
                    "Stabilizing global IT operations.",
                    "Unlocking population-health insights in healthcare.",
                    "Cost efficiency and operational excellence in cloud.",
                ],
            },
            {
                "heading": "Events, Webinars & News",
                "bullets": [
                    "Events such as the ET GCC Conclave and Microsoft AI Tour.",
                    "Webinars on resilient cloud infrastructure, API testing and product demos.",
                    "News including Vector 2 launch and company milestones.",
                ],
            },
            {
                "heading": "Read More",
                "paras": [
                    "Browse the full library at parkar.in/insights.",
                ],
            },
        ],
    },
}


if __name__ == "__main__":
    files = build_all()
    print(f"Generated {len(files)} documents in {DOCX_DIR}:")
    for name in files:
        print(f"  - {name}")
